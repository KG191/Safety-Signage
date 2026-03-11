#!/usr/bin/env python3
"""Generate the ML training methodology report as a Word document."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# --- Title ---
title = doc.add_heading('ML Sign Detection Model — Training Methodology Report', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

meta = doc.add_paragraph()
meta.style.font.size = Pt(10)
meta.add_run('Safety Signage Audit Application — AS 1319-1994 Compliance Tool\n').bold = True
meta.add_run('Date: March 2026  |  Author: Kiran Gokal, WSP\n')
meta.add_run('Classification: Internal / Technical Reference')

# --- 1. Introduction ---
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'This report describes the machine learning methodology used to develop an automated '
    'safety sign category classifier for the Safety Signage Audit application. The classifier '
    'identifies signs conforming to AS 1319-1994 (Safety signs for the occupational environment) '
    'and operates as part of a multi-layered detection pipeline that includes local computer vision '
    'analysis and optional cloud-based Vision AI.'
)
doc.add_paragraph(
    'The primary objective was to replace a generic ImageNet-based MobileNet model — which achieved '
    'approximately 25% category accuracy on safety signs — with a purpose-trained model capable of '
    'classifying signs into the seven AS 1319 categories with field-usable accuracy (target: >85%).'
)

# --- 2. Dataset Selection ---
doc.add_heading('2. Dataset Selection and Rationale', level=1)
doc.add_paragraph(
    'No publicly available machine learning dataset exists for Australian AS 1319 safety signs. '
    'A review of available resources identified the ORP-SIG-2024 dataset as the most suitable '
    'training source.'
)

doc.add_heading('2.1 ORP-SIG-2024 Dataset', level=2)
doc.add_paragraph(
    'The ORP-SIG-2024 dataset (Mendeley Data, DOI: 10.17632/dfg5hnxrzg.1) contains 94,080 images '
    'across 299 ISO 7010 pictograms organised into five categories: Prohibition (P-series), '
    'Mandatory (M-series), Warning (W-series), Emergency/Safe Condition (E-series), and '
    'Fire Protection (F-series). Each pictogram includes original and augmented variants with '
    'transformations such as rotation, brightness variation, blur, and cropping to simulate '
    'real-world degradation.'
)
doc.add_paragraph(
    'Published validation on this dataset reports 99.98% accuracy on clean pictograms and '
    '95.30% on degraded images (Frontiers in Public Health, 2024).'
)

doc.add_heading('2.2 ISO 7010 to AS 1319 Mapping', level=2)
doc.add_paragraph(
    'ISO 7010 (Graphical symbols — Safety colours and safety signs) and AS 1319-1994 share a '
    'common visual framework derived from earlier ISO standards. Both standards use identical '
    'safety colours (red, blue, yellow, green), identical geometric shapes for each category, '
    'and substantially overlapping pictogram designs. The category-level mapping is direct:'
)

# Mapping table
table = doc.add_table(rows=8, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.LEFT
headers = ['ISO 7010 Series', 'AS 1319 Category', 'Shape / Colour', 'Mapping Confidence']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

data = [
    ('P-series', 'Prohibition', 'Red circle + diagonal bar', 'Direct'),
    ('M-series', 'Mandatory', 'Blue filled circle', 'Direct'),
    ('W-series', 'Warning', 'Yellow triangle', 'Direct'),
    ('E-series', 'Emergency Information', 'Green rectangle', 'Direct'),
    ('F-series', 'Fire', 'Red rectangle', 'Direct'),
    ('—', 'Danger (AS 1319 specific)', 'Red oval + black bg', 'Synthesised'),
    ('—', 'Restriction (AS 1319 specific)', 'Red circle, no bar', 'Synthesised'),
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r + 1].cells[c].text = val

doc.add_paragraph()  # spacer

doc.add_paragraph(
    'Two AS 1319 categories have no ISO 7010 equivalent. DANGER signs (Clause 2.3.4) use a '
    'distinctive red oval with white "DANGER" text on a black background — these are word-only '
    'signs unique to the Australian standard. Restriction signs (red annulus, white interior, '
    'no diagonal bar) represent a subset that ISO 7010 does not distinguish from prohibition. '
    'Training images for both categories were synthetically generated (500 images each) with '
    'randomised parameters including text content, rotation, brightness, and noise to provide '
    'visual diversity.'
)

# --- 3. Model Architecture ---
doc.add_heading('3. Model Architecture and Training', level=1)

doc.add_heading('3.1 Architecture', level=2)
doc.add_paragraph(
    'The classifier uses MobileNet v3 Small as a feature extraction backbone with a custom '
    'classification head. MobileNet v3 Small was selected for its balance of accuracy and '
    'inference speed on mobile devices — the model runs in-browser via TensorFlow.js at under '
    '500ms on mid-range smartphones.'
)

p = doc.add_paragraph()
p.add_run('Model structure:').bold = True
doc.add_paragraph('MobileNet v3 Small (ImageNet weights, 939K parameters) — frozen', style='List Bullet')
doc.add_paragraph('Global Average Pooling 2D', style='List Bullet')
doc.add_paragraph('Dense(128, ReLU)', style='List Bullet')
doc.add_paragraph('Dropout(0.3)', style='List Bullet')
doc.add_paragraph('Dense(7, softmax) — outputs 7 AS 1319 category probabilities', style='List Bullet')

doc.add_paragraph(
    'Input preprocessing normalises images to 224×224 pixels with pixel values scaled to [-1, 1], '
    'matching MobileNet v3 expectations. The total model size after uint16 quantization is 2.1 MB.'
)

doc.add_heading('3.2 Training Process', level=2)
doc.add_paragraph(
    'Training used the organised dataset of 6,083 images (5,083 from ORP-SIG-2024 plus 1,000 '
    'synthesised) split 80/10/10 into train/validation/test sets. Data augmentation during '
    'training included random rotation (±29°), zoom (±10%), translation (±5%), brightness '
    'variation (±20%), contrast adjustment (±20%), and horizontal flipping.'
)

p = doc.add_paragraph()
p.add_run('Training configuration:').bold = True
doc.add_paragraph('Optimiser: Adam, learning rate 1×10⁻³ (frozen base) then 1×10⁻⁴ (unfrozen)', style='List Bullet')
doc.add_paragraph('Loss: Sparse categorical cross-entropy', style='List Bullet')
doc.add_paragraph('Batch size: 32', style='List Bullet')
doc.add_paragraph('Early stopping: patience 5, restoring best weights on validation accuracy', style='List Bullet')
doc.add_paragraph('Learning rate reduction: factor 0.5 on validation loss plateau (patience 3)', style='List Bullet')

doc.add_paragraph(
    'The final deployed model used a two-phase approach: Phase 1 trained only the classification '
    'head (25 epochs with frozen convolutional base), followed by unfreezing the top 10 '
    'convolutional layers for an additional 25 epochs at a reduced learning rate. This achieved '
    '89.5% test accuracy on the held-out test set.'
)

# --- 4. Results ---
doc.add_heading('4. Results', level=1)

# Results table
table2 = doc.add_table(rows=9, cols=5)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.LEFT
headers2 = ['Category', 'Precision', 'Recall', 'F1-Score', 'Test Samples']
for i, h in enumerate(headers2):
    table2.rows[0].cells[i].text = h
    table2.rows[0].cells[i].paragraphs[0].runs[0].bold = True

results = [
    ('Prohibition', '1.00', '0.83', '0.91', '118'),
    ('Mandatory', '0.69', '0.91', '0.79', '92'),
    ('Restriction*', '1.00', '0.96', '0.98', '53'),
    ('Warning', '0.92', '0.99', '0.96', '145'),
    ('Danger*', '1.00', '1.00', '1.00', '41'),
    ('Emergency', '0.90', '0.93', '0.91', '121'),
    ('Fire', '1.00', '0.42', '0.59', '38'),
    ('Overall', '0.92', '0.90', '0.90', '608'),
]
for r, row_data in enumerate(results):
    for c, val in enumerate(row_data):
        cell = table2.rows[r + 1].cells[c]
        cell.text = val
        if r == 7:  # Overall row bold
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()
doc.add_paragraph(
    '* Restriction and Danger categories used synthesised training data. '
    'The high scores reflect the visual distinctiveness of synthesised images and should be '
    'validated against photographs of real signs in the field.'
)
doc.add_paragraph(
    'Fire category recall (42%) is the weakest result, attributable to the small number of '
    'fire protection pictograms in the source dataset (19 originals, 323 total with augmentation) '
    'and visual similarity to prohibition signs (both use red). Additional field photographs of '
    'fire safety signs would improve this category.'
)

# --- 5. Integration ---
doc.add_heading('5. Integration and Deployment', level=1)
doc.add_paragraph(
    'The trained model was converted to TensorFlow.js GraphModel format with uint16 quantization '
    '(2.1 MB total) and is served alongside the application via GitHub Pages. The model loads '
    'lazily when the user enters the Capture view and is cached by the browser for offline use.'
)
doc.add_paragraph(
    'Within the detection pipeline, the ML classifier operates alongside colour analysis and '
    'shape detection. The ML prediction receives 40% weight in the composite confidence score '
    '(increased from 20% with the previous ImageNet model) and is used to resolve ambiguous '
    'cases — particularly distinguishing prohibition from restriction signs (both red circles) '
    'and danger from fire signs (both red rectangles). The Vision API path, when configured, '
    'takes priority over both local methods.'
)

# --- 6. Limitations ---
doc.add_heading('6. Limitations and Future Work', level=1)

doc.add_paragraph(
    'Training data consists of standardised pictogram images, not photographs of installed signs. '
    'Real-world factors — weathering, fading, partial obstruction, non-standard printing, and '
    'varied lighting — are only partially represented by the dataset augmentations. Field '
    'validation with photographs from Australian sites is essential before relying on ML '
    'predictions for compliance decisions.',
)

doc.add_paragraph(
    'The model classifies at the category level only (Phase 1). Specific sign number '
    'identification (e.g., distinguishing sign 424 "Head protection" from sign 425 '
    '"Hearing protection") requires a Phase 2 model trained on individual pictograms, '
    'which would map the 299 ISO 7010 pictograms to the ~30 AS 1319 Appendix B sign numbers.'
)

doc.add_paragraph(
    'AS 1319-1994 was reconfirmed in 2018 but has not been revised to align with the current '
    'ISO 7010:2019 edition. Some newer ISO 7010 pictograms (added post-2003) may not have '
    'AS 1319 equivalents. The mapping used in this work covers the core categories that have '
    'remained stable across both standards.'
)

# --- References ---
doc.add_heading('References', level=1)
refs = [
    'AS 1319-1994 (Reconfirmed 2018), Safety signs for the occupational environment. Standards Australia.',
    'ISO 7010:2019, Graphical symbols — Safety colours and safety signs — Registered safety signs. '
    'International Organization for Standardization.',
    'ORP-SIG-2024 Dataset. Mendeley Data, DOI: 10.17632/dfg5hnxrzg.1.',
    'Fernández-Villacañas Marín, M. et al. (2024). "Machine vision-based recognition methodology '
    'of standard safety signs in work environments." Frontiers in Public Health, 12, 1431757.',
    'Howard, A. et al. (2019). "Searching for MobileNetV3." Proceedings of the IEEE/CVF '
    'International Conference on Computer Vision.',
]
for i, ref in enumerate(refs, 1):
    doc.add_paragraph(f'[{i}] {ref}', style='List Number')

# Save
output_path = '/Users/kirangokal/Documents/SafetySignages/ML_Training_Methodology_Report.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
